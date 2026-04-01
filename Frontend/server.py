import sys
import json
import logging
from pathlib import Path
import io
import re
from datetime import datetime


LOG_DIR = Path(__file__).resolve().parent / "logs"
LOG_DIR.mkdir(exist_ok=True)
LOG_FILE = LOG_DIR / "app.log"


class ImmediateFileHandler(logging.FileHandler):
    """File handler that flushes after every write."""

    def emit(self, record):
        super().emit(record)
        self.flush()


class ImmediateStreamHandler(logging.StreamHandler):
    """Stream handler that flushes after every write."""

    def emit(self, record):
        super().emit(record)
        self.flush()


class PrintToLogger:
    """Redirects print() to both console and log file immediately."""

    def __init__(self, original, logger, level=logging.INFO):
        self.original = original
        self.logger = logger
        self.level = level
        self._buffer = ""

    def write(self, text):
        if self.original:
            self.original.write(text)
            self.original.flush()
        self._buffer += text
        while "\n" in self._buffer:
            line, self._buffer = self._buffer.split("\n", 1)
            line = line.rstrip("\r")
            if line.strip():
                self.logger.log(self.level, line)

    def flush(self):
        # Flush remaining buffer
        if self._buffer.strip():
            self.logger.log(self.level, self._buffer.strip())
            self._buffer = ""
        if self.original:
            self.original.flush()

    def isatty(self):
        return hasattr(self.original, "isatty") and self.original.isatty()


def setup_logging():
    """Configure ALL logging to single file + console with IMMEDIATE flush."""

    # ──── File handler ────
    file_handler = ImmediateFileHandler(LOG_FILE, mode="w", encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(
        logging.Formatter(
            "%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
            datefmt="%Y-%m-%d %H:%M:%S",
        )
    )

    console_handler = ImmediateStreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(
        logging.Formatter(
            "%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
            datefmt="%H:%M:%S",
        )
    )

    root = logging.getLogger()
    root.setLevel(logging.DEBUG)
    root.handlers.clear()
    root.addHandler(file_handler)
    root.addHandler(console_handler)

    loggers_to_capture = [
        "uvicorn",
        "uvicorn.error",
        "uvicorn.access",
        "fastapi",
        "httpx",
        "httpcore",
        "langchain",
        "langchain_core",
        "langchain_mistralai",
        "chromadb",
        "PRINT",
        "STDERR",
        # Our own modules - use propagate=True so they flow to root
        "Backend",
        "Backend.Services",
        "Backend.Services.Agent_Utils",
        "Backend.Services.GraphNodes",
        "Backend.Services.ResearchGraph",
    ]

    for logger_name in loggers_to_capture:
        lib_logger = logging.getLogger(logger_name)
        lib_logger.handlers.clear()
        lib_logger.addHandler(file_handler)
        lib_logger.addHandler(console_handler)
        lib_logger.setLevel(logging.DEBUG)
        lib_logger.propagate = False

    # ──── Redirect print() and stderr ────
    print_logger = logging.getLogger("PRINT")
    stderr_logger = logging.getLogger("STDERR")
    sys.stdout = PrintToLogger(sys.__stdout__, print_logger, logging.INFO)
    sys.stderr = PrintToLogger(sys.__stderr__, stderr_logger, logging.ERROR)

    # ──── Startup banner ────
    root.info("=" * 60)
    root.info("LOGGING INITIALIZED")
    root.info(f"LOG FILE: {LOG_FILE.absolute()}")
    root.info("=" * 60)

    return file_handler, console_handler


# ──── Initialize logging FIRST ────
_file_handler, _console_handler = setup_logging()
logger = logging.getLogger(__name__)


def log_flush():
    """Force flush all handlers."""
    _file_handler.flush()
    _console_handler.flush()


# ──────────────────────────────────────────────────────────────
# NOW IMPORT EVERYTHING ELSE
# ──────────────────────────────────────────────────────────────

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

logger.info("Importing Backend modules...")
log_flush()

from fastapi import FastAPI, UploadFile, File, Request
from fastapi.responses import StreamingResponse, HTMLResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from starlette.middleware.base import BaseHTTPMiddleware
from contextlib import asynccontextmanager

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import markdown

from Backend.Services.Agent_Utils import Agent

logger.info("All imports complete.")
log_flush()


# ──────────────────────────────────────────────────────────────
# SESSION STATE  (defined before lifespan so it's available)
# ──────────────────────────────────────────────────────────────

session = {
    "model": "mistral-medium-latest",
    "mode": "General",
    "current_agent_mode": None,
    "current_agent_model": None,
    "agent": None,
    "file_stored": False,
    "file_type": "text",
    "uploaded_filename": None,
}


# ──────────────────────────────────────────────────────────────
# LIFESPAN
# ──────────────────────────────────────────────────────────────


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Runs on startup and shutdown."""
    logger.info("=" * 60)
    logger.info("SERVER STARTING UP")
    logger.info(f"Mode : {session['mode']}")
    logger.info(f"Model: {session['model']}")
    logger.info("=" * 60)
    log_flush()

    yield

    logger.info("=" * 60)
    logger.info("SERVER SHUTTING DOWN")
    logger.info("=" * 60)
    log_flush()


# ──────────────────────────────────────────────────────────────
# APP + MIDDLEWARE
# ──────────────────────────────────────────────────────────────

app = FastAPI(title="Research AI - Jarvis", lifespan=lifespan)

# Mount static files (favicon, assets, etc.)
_STATIC_DIR = Path(__file__).parent / "static"
_STATIC_DIR.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(_STATIC_DIR)), name="static")


class RequestLoggingMiddleware(BaseHTTPMiddleware):
    """Logs every incoming request and outgoing response status."""

    async def dispatch(self, request, call_next):
        logger.info(f">>> {request.method} {request.url.path}")
        log_flush()
        try:
            response = await call_next(request)
            logger.info(
                f"<<< {request.method} {request.url.path} -> {response.status_code}"
            )
            log_flush()
            return response
        except Exception as e:
            logger.error(
                f"!!! {request.method} {request.url.path} FAILED: {e}",
                exc_info=True,
            )
            log_flush()
            raise


app.add_middleware(RequestLoggingMiddleware)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://researcher-frontend-eta.vercel.app",
        "http://localhost:3000",
        "http://localhost:8000",
        "http://127.0.0.1:5500",
        "http://127.0.0.1:8000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ──────────────────────────────────────────────────────────────
# AGENT FACTORY
# ──────────────────────────────────────────────────────────────


def get_or_create_agent() -> Agent:
    """Create agent if needed, or recreate on mode/model change."""
    mode = session["mode"]
    model = session["model"]
    has_document = session.get("file_stored", False)

    needs_rebuild = (
        session["agent"] is None
        or session["current_agent_mode"] != mode
        or session["current_agent_model"] != model
    )

    if needs_rebuild:
        logger.info(f"[Agent Factory] Creating Agent: model={model}, role={mode}, has_document={has_document}")
        log_flush()
        try:
            session["agent"] = Agent(
                model_name=model,
                agent_role=mode,
                has_document=has_document,
            )
            session["current_agent_mode"] = mode
            session["current_agent_model"] = model
            logger.info("[Agent Factory] Agent created successfully")
            log_flush()
        except Exception as e:
            logger.error(f"[Agent Factory] Failed to create Agent: {e}", exc_info=True)
            log_flush()
            raise

    return session["agent"]


# ──────────────────────────────────────────────────────────────
# ENDPOINTS
# ──────────────────────────────────────────────────────────────


@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    logger.debug("Serving frontend HTML")
    log_flush()
    html_path = Path(__file__).parent / "index.html"
    return FileResponse(html_path, media_type="text/html")


@app.post("/api/config")
async def update_config(request: Request):
    """Update model and mode settings."""
    data = await request.json()
    logger.info(f"[/api/config] Received: {data}")

    if "model" in data:
        session["model"] = data["model"]
    if "mode" in data:
        session["mode"] = data["mode"]
        if data["mode"] == "Research":
            session["model"] = "mistral-medium-latest"

    logger.info(
        f"[/api/config] Updated → model={session['model']}, mode={session['mode']}"
    )
    log_flush()

    return {
        "status": "ok",
        "model": session["model"],
        "mode": session["mode"],
    }


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """Upload a PDF/image file to use as context."""
    logger.info(f"[/api/upload] File: {file.filename} | Type: {file.content_type}")
    log_flush()

    agent = get_or_create_agent()
    content = await file.read()
    content_type = file.content_type or ""
    filename = file.filename or "unknown"

    f_type = "text"
    if "pdf" in content_type:
        f_type = "pdf"
        try:
            logger.info(f"[/api/upload] Processing PDF: {filename}")
            log_flush()
            agent.convert_and_store_to_vect_db(content)
            session["file_stored"] = True
            logger.info("[/api/upload] PDF stored successfully")
            log_flush()
        except ValueError as e:
            logger.error(f"[/api/upload] PDF processing failed: {e}")
            log_flush()
            return {"status": "error", "message": str(e)}
        except Exception as e:
            logger.error(f"[/api/upload] Upload failed: {e}", exc_info=True)
            log_flush()
            return {"status": "error", "message": f"Failed to process PDF: {str(e)}"}
    elif "image" in content_type:
        f_type = "image"
        logger.info("[/api/upload] Image received (not stored in vector DB)")
        log_flush()

    session["file_type"] = f_type
    session["uploaded_filename"] = filename

    return {
        "status": "ok",
        "filename": filename,
        "file_type": f_type,
        "stored": session["file_stored"],
    }


@app.post("/api/clear-file")
async def clear_file():
    """Remove uploaded file and clear vectorstore."""
    logger.info("[/api/clear-file] Clearing file and vectorstore")
    log_flush()

    if session["agent"] and session["file_stored"]:
        try:
            session["agent"].clear_vectorstore()
            logger.info("[/api/clear-file] Vectorstore cleared")
            log_flush()
        except Exception as e:
            logger.error(f"[/api/clear-file] Clear vectorstore failed: {e}")
            log_flush()

    session["file_stored"] = False
    session["file_type"] = "text"
    session["uploaded_filename"] = None
    return {"status": "ok"}


@app.post("/api/chat")
async def chat(request: Request):
    """Stream agent response via SSE."""
    data = await request.json()
    message = data.get("message", "").strip()

    logger.info("=" * 40)
    logger.info("[/api/chat] CHAT REQUEST RECEIVED")
    logger.info(
        f"[/api/chat] Message : '{message[:120]}{'...' if len(message) > 120 else ''}'"
    )
    logger.info(f"[/api/chat] Mode    : {session['mode']}")
    logger.info(f"[/api/chat] FileType: {session['file_type']}")
    logger.info("=" * 40)
    log_flush()

    if not message:
        logger.warning("[/api/chat] Empty message received")
        log_flush()

        def _empty_error():
            yield f"data: {json.dumps({'type': 'error', 'content': 'Empty message'})}\n\n"

        return StreamingResponse(_empty_error(), media_type="text/event-stream")

    try:
        agent = get_or_create_agent()
    except Exception as e:
        logger.error(f"[/api/chat] Failed to get agent: {e}", exc_info=True)
        log_flush()

        def _agent_error():
            yield f"data: {json.dumps({'type': 'error', 'content': str("Error")})}\n\n"  # noqa: F821

        return StreamingResponse(_agent_error(), media_type="text/event-stream")

    f_type = session["file_type"]
    logger.info(f"[/api/chat] Invoking agent with file_type={f_type}")
    log_flush()

    def generate():
        chunk_count = 0
        total_chars = 0
        try:
            logger.info("[/api/chat] Stream generation started")
            log_flush()

            stream = agent.Invoke_agent(message, f_type)
            for chunk in stream:
                chunk_count += 1
                total_chars += len(chunk)
                payload = json.dumps({"type": "chunk", "content": chunk})
                yield f"data: {payload}\n\n"

                # Log progress every 10 chunks
                if chunk_count % 10 == 0:
                    logger.debug(
                        f"[/api/chat] Progress: {chunk_count} chunks | {total_chars} chars"
                    )
                    log_flush()

            logger.info(
                f"[/api/chat] COMPLETE: {chunk_count} chunks | {total_chars} chars"
            )
            log_flush()
            yield f"data: {json.dumps({'type': 'done'})}\n\n"

        except Exception as e:
            logger.error(f"[/api/chat] Stream error: {e}", exc_info=True)
            log_flush()
            yield f"data: {json.dumps({'type': 'error', 'content': str(e)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/clear")
async def clear_conversation():
    """Clear conversation — reset agent to force new memory."""
    logger.info("[/api/clear] Clearing conversation and session")
    log_flush()

    if session["agent"] and session["file_stored"]:
        try:
            session["agent"].clear_vectorstore()
            logger.info("[/api/clear] Vectorstore cleared")
            log_flush()
        except Exception as e:
            logger.error(f"[/api/clear] Vectorstore clear failed: {e}")
            log_flush()

    session["agent"] = None
    session["current_agent_mode"] = None
    session["current_agent_model"] = None
    session["file_stored"] = False
    session["file_type"] = "text"
    session["uploaded_filename"] = None

    logger.info("[/api/clear] Session cleared")
    log_flush()
    return {"status": "ok"}


@app.get("/api/status")
async def status():
    logger.debug("[/api/status] Status check")
    log_flush()
    return {
        "status": "ready",
        "model": session["model"],
        "mode": session["mode"],
        "file_stored": session["file_stored"],
        "uploaded_filename": session["uploaded_filename"],
    }


@app.post("/api/download")
async def download_response(request: Request):
    """
    Generate downloadable Word or PDF file from response text.
    Expects JSON: {"text": "response text", "format": "word" or "pdf"}
    """
    logger.info("[/api/download] Download request received")
    log_flush()
    
    try:
        data = await request.json()
        response_text = data.get("text", "")
        file_format = data.get("format", "word").lower()
        
        if not response_text:
            return {"status": "error", "message": "No text provided"}
        
        # Clean markdown syntax for better formatting
        cleaned_text = response_text
        
        # Generate timestamp for filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if file_format == "word":
            # Generate Word document
            logger.info("[/api/download] Generating Word document...")
            log_flush()
            
            doc = Document()
            
            # Add title
            title = doc.add_heading('Research Response', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.size = Pt(10)
            date_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()  # Blank line
            
            # Process text - handle markdown headings and citations
            lines = cleaned_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                # Handle markdown headings
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
            
            # Save to bytes buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            filename = f"research_response_{timestamp}.docx"
            
            logger.info(f"[/api/download] Word document generated: {filename}")
            log_flush()
            
            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        
        elif file_format == "pdf":
            # Generate PDF document
            logger.info("[/api/download] Generating PDF document...")
            log_flush()
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Title
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "Research Response", ln=True, align="C")
            pdf.ln(5)
            
            # Timestamp
            pdf.set_font("Arial", "I", 10)
            pdf.set_text_color(128, 128, 128)
            pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align="C")
            pdf.ln(5)
            
            # Reset color
            pdf.set_text_color(0, 0, 0)
            
            # Process text
            lines = cleaned_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    pdf.ln(5)
                    continue
                
                # Clean line for PDF - handle unicode early
                line_clean = line.encode('latin-1', 'ignore').decode('latin-1')
                
                # Handle markdown headings
                if line_clean.startswith('### '):
                    pdf.ln(2)
                    pdf.set_font("Arial", "B", 12)
                    pdf.multi_cell(0, 8, line_clean[4:])
                    pdf.ln(2)
                    pdf.set_font("Arial", "", 11)
                elif line_clean.startswith('## '):
                    pdf.set_font("Arial", "B", 14)
                    pdf.multi_cell(0, 9, line_clean[3:])
                    pdf.ln(3)
                    pdf.set_font("Arial", "", 11)
                elif line_clean.startswith('# '):
                    pdf.set_font("Arial", "B", 16)
                    pdf.multi_cell(0, 10, line_clean[2:])
                    pdf.ln(4)
                    pdf.set_font("Arial", "", 11)
                else:
                    # Regular text - remove markdown bold syntax
                    line_clean = line_clean.replace('**', '')
                    pdf.set_font("Arial", "", 11)
                    try:
                        pdf.multi_cell(0, 6, line_clean)
                    except Exception as e:
                        logger.warning(f"[/api/download] Skipped line: {str(e)[:50]}")
            
            # Save to bytes buffer - fpdf2's output(dest='S') returns bytes directly
            buffer = io.BytesIO()
            pdf_bytes = pdf.output(dest='S')
            # Handle both old and new fpdf2 versions
            if isinstance(pdf_bytes, bytes):
                buffer.write(pdf_bytes)
            else:
                # Older versions might return str
                buffer.write(pdf_bytes.encode('latin-1', 'ignore'))
            buffer.seek(0)
            
            filename = f"research_response_{timestamp}.pdf"
            
            logger.info(f"[/api/download] PDF document generated: {filename}")
            log_flush()
            
            return StreamingResponse(
                buffer,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        
        else:
            return {"status": "error", "message": f"Unsupported format: {file_format}"}
    
    except Exception as e:
        logger.error(f"[/api/download] Error: {e}", exc_info=True)
        log_flush()
        return {"status": "error", "message": str(e)}


# ──────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn

    logger.info("Starting uvicorn server...")
    log_flush()

    # Custom uvicorn logging config to use our handlers
    uvicorn_log_config = {
        "version": 1,
        "disable_existing_loggers": False,
        "formatters": {
            "default": {
                "format": "%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
                "datefmt": "%H:%M:%S",
            },
        },
        "handlers": {
            "default": {
                "formatter": "default",
                "class": "logging.StreamHandler",
                "stream": "ext://sys.stdout",
            },
        },
        "loggers": {
            "uvicorn": {"handlers": ["default"], "level": "INFO", "propagate": False},
            "uvicorn.error": {"level": "INFO"},
            "uvicorn.access": {
                "handlers": ["default"],
                "level": "INFO",
                "propagate": False,
            },
        },
    }

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000,
        log_level="info",
        log_config=uvicorn_log_config,
    )
